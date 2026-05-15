"""
Extractor de texto de documentos — Sprint 6.

Toma un `Item` con `media_tipo='documento'` (con `media.Attachment` asociado),
descarga el binario de MinIO, detecta el formato y extrae texto:

  - PDF: pdfplumber (mantiene mejor el layout que pypdf)
  - DOCX: python-docx (parrafos + tablas)
  - XLSX: openpyxl (todas las hojas, fila por fila)
  - TXT/CSV/MD: chardet para detectar encoding, después decode

El texto extraído va a `item.contenido` (si estaba vacío) y la metadata a
`item.datos['extraccion']` (formato, páginas/hojas, chars, modelo, etc.).
Después de extraer encadena con `encolar_job_embed` para que entre al chat.
"""

from datetime import datetime, timezone
from io import BytesIO

import chardet
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.config import get_settings
from app.core.logging import get_logger
from app.models.core import Item
from app.models.media import Attachment
from app.models.processing import Job
from app.services.embedder import encolar_job_embed
from app.services.minio_client import VaultStorage

logger = get_logger(__name__)
settings = get_settings()

# Máximo de caracteres que guardamos en item.contenido (para no inflar Postgres)
_MAX_CONTENIDO = 50_000
# Si un docx/pdf tiene más páginas/hojas, sólo extraemos las primeras
_MAX_PAGES = 200


# ---------------------------------------------------------------------------
# Cola (processing.jobs)
# ---------------------------------------------------------------------------


def encolar_job_extract(db: Session, item_id) -> bool:
    """Crea un Job pendiente para extraer texto del documento. Idempotente."""
    if item_id is None:
        return False
    existente = db.execute(
        select(Job.id).where(
            Job.item_id == item_id,
            Job.tipo == "extract",
            Job.estado.in_(["pendiente", "en_proceso"]),
        )
    ).first()
    if existente:
        return False
    db.add(Job(tipo="extract", item_id=item_id, estado="pendiente"))
    return True


# ---------------------------------------------------------------------------
# Extractores por formato
# ---------------------------------------------------------------------------


def _ext_de_filename(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower().strip()


def _extraer_pdf(content: bytes) -> dict:
    import pdfplumber

    out_paginas: list[str] = []
    metadata: dict = {}
    total_pages = 0
    with pdfplumber.open(BytesIO(content)) as pdf:
        metadata = dict(pdf.metadata or {})
        total_pages = len(pdf.pages)
        for p in pdf.pages[:_MAX_PAGES]:
            try:
                txt = p.extract_text() or ""
            except Exception:  # noqa: BLE001
                txt = ""
            out_paginas.append(txt.strip())
    return {
        "formato": "pdf",
        "texto": "\n\n".join(t for t in out_paginas if t),
        "paginas_extraidas": len(out_paginas),
        "paginas_totales": total_pages,
        "metadata_pdf": {k: str(v)[:200] for k, v in metadata.items() if v is not None},
    }


def _extraer_docx(content: bytes) -> dict:
    from docx import Document

    doc = Document(BytesIO(content))
    parrafos = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
    # Tablas: aplastadas como filas de texto separadas por tab
    tablas_txt: list[str] = []
    for tbl in doc.tables:
        for row in tbl.rows:
            celdas = [c.text.strip() for c in row.cells]
            tablas_txt.append("\t".join(celdas))
    texto = "\n".join(parrafos)
    if tablas_txt:
        texto = (texto + "\n\n[TABLAS]\n" + "\n".join(tablas_txt)).strip()
    return {
        "formato": "docx",
        "texto": texto,
        "n_parrafos": len(parrafos),
        "n_tablas": len(doc.tables),
    }


def _extraer_xlsx(content: bytes) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    hojas_out: list[str] = []
    for nombre in wb.sheetnames:
        ws = wb[nombre]
        filas_txt: list[str] = []
        for row in ws.iter_rows(values_only=True):
            celdas = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in celdas):
                filas_txt.append("\t".join(celdas))
        if filas_txt:
            hojas_out.append(f"[HOJA: {nombre}]\n" + "\n".join(filas_txt))
    wb.close()
    return {
        "formato": "xlsx",
        "texto": "\n\n".join(hojas_out),
        "n_hojas": len(wb.sheetnames),
        "hojas": list(wb.sheetnames),
    }


def _extraer_texto_plano(content: bytes, formato: str) -> dict:
    # Detección de encoding
    sample = content[:200_000]  # con 200KB suele alcanzar
    enc = (chardet.detect(sample).get("encoding") or "utf-8").lower()
    try:
        texto = content.decode(enc, errors="replace")
    except (LookupError, UnicodeDecodeError):
        texto = content.decode("utf-8", errors="replace")
    return {"formato": formato, "texto": texto, "encoding_detectado": enc}


def _detectar_y_extraer(content: bytes, filename: str | None, mime: str | None) -> dict:
    """Despacho por extensión / MIME. Devuelve dict con 'formato', 'texto', extras."""
    ext = _ext_de_filename(filename)
    mime = (mime or "").lower()

    # PDF
    if ext == "pdf" or "pdf" in mime or content[:4] == b"%PDF":
        return _extraer_pdf(content)
    # DOCX (zip que empieza con PK)
    if ext == "docx" or "wordprocessingml" in mime:
        return _extraer_docx(content)
    # XLSX
    if ext in ("xlsx", "xlsm") or "spreadsheetml" in mime:
        return _extraer_xlsx(content)
    # Texto plano
    if ext in ("txt", "md", "csv", "log") or mime.startswith("text/"):
        return _extraer_texto_plano(content, ext or "txt")

    raise ValueError(f"formato no soportado (ext={ext}, mime={mime})")


# ---------------------------------------------------------------------------
# Núcleo: extraer de un item
# ---------------------------------------------------------------------------


def _bucket_y_key(minio_path: str) -> tuple[str, str]:
    bucket, _, key = minio_path.partition("/")
    return bucket or settings.minio_bucket_raw, key


def extraer_item(db: Session, item: Item, *, vault: VaultStorage | None = None) -> dict:
    """Extrae texto del documento de un item. Idempotente: re-extrae si lo llamás de
    nuevo. No commitea — el caller maneja la transacción.
    """
    if item.media_tipo != "documento":
        return {"ok": False, "error": f"item no es documento (media_tipo={item.media_tipo})"}

    att = db.execute(
        select(Attachment).where(Attachment.item_id == item.id, Attachment.tipo == "documento")
    ).scalars().first()
    if att is None:
        return {"ok": False, "error": "sin attachment de documento"}

    vault = vault or VaultStorage()
    bucket, key = _bucket_y_key(att.minio_path)
    try:
        content = vault.get(bucket, key)
    except Exception as e:  # noqa: BLE001
        return {"ok": False, "error": f"minio: {e}"}
    if not content:
        return {"ok": False, "error": "binario vacío"}

    try:
        res = _detectar_y_extraer(content, att.filename_original, att.mime_type)
    except Exception as e:  # noqa: BLE001
        return {"ok": False, "error": f"extract: {e}"}

    texto = (res.get("texto") or "").strip()
    ahora = datetime.now(timezone.utc).isoformat()

    if texto and not (item.contenido or "").strip():
        item.contenido = texto[:_MAX_CONTENIDO]

    extras = {k: v for k, v in res.items() if k != "texto"}
    nuevos_datos = dict(item.datos or {})
    nuevos_datos["extraccion"] = {
        **extras,
        "chars": len(texto),
        "truncado": len(texto) > _MAX_CONTENIDO,
        "extracted_at": ahora,
    }
    item.datos = nuevos_datos

    att.procesado = True
    att.nivel_procesamiento = 1
    att_datos = dict(att.datos or {})
    att_datos["extracted_at"] = ahora
    att.datos = att_datos

    return {
        "ok": True,
        "item_id": str(item.id),
        "formato": res.get("formato"),
        "chars": len(texto),
        "extras": extras,
    }


# ---------------------------------------------------------------------------
# Worker
# ---------------------------------------------------------------------------


def procesar_jobs(db: Session, limit: int = 20) -> dict:
    """Drena hasta `limit` jobs pendientes de tipo 'extract'."""
    limit = max(1, min(limit, 200))
    vault = VaultStorage()

    pendientes = db.execute(
        select(Job)
        .where(Job.tipo == "extract", Job.estado == "pendiente")
        .order_by(Job.created_at.asc())
        .limit(limit)
    ).scalars().all()

    procesados = exitosos = fallidos = 0
    errores: list[str] = []

    for job in pendientes:
        job.estado = "en_proceso"
        job.started_at = datetime.now(timezone.utc)
        job.intentos = (job.intentos or 0) + 1
        db.commit()

        try:
            item = db.get(Item, job.item_id) if job.item_id else None
            if item is None:
                job.estado = "fallido"
                job.error = "item no encontrado"
                job.completed_at = datetime.now(timezone.utc)
                db.commit()
                fallidos += 1
                procesados += 1
                continue

            res = extraer_item(db, item, vault=vault)
            if not res["ok"]:
                raise RuntimeError(res.get("error") or "extract falló")

            job.estado = "completado"
            job.resultado = res
            job.completed_at = datetime.now(timezone.utc)
            encolar_job_embed(db, item.id)
            db.commit()
            exitosos += 1
        except Exception as e:  # noqa: BLE001
            db.rollback()
            j2 = db.get(Job, job.id)
            if j2 is not None:
                if (j2.intentos or 0) >= (j2.max_intentos or 3):
                    j2.estado = "fallido"
                    j2.completed_at = datetime.now(timezone.utc)
                else:
                    j2.estado = "pendiente"
                j2.error = str(e)[:1000]
                db.commit()
            errores.append(f"{job.id}: {str(e)[:200]}")
            fallidos += 1
            logger.error("extract_job_failed", job_id=str(job.id), error=str(e))
        procesados += 1

    pendientes_restantes = db.execute(
        select(func.count()).select_from(Job).where(Job.tipo == "extract", Job.estado == "pendiente")
    ).scalar_one()
    return {
        "procesados": procesados,
        "exitosos": exitosos,
        "fallidos": fallidos,
        "pendientes_restantes": int(pendientes_restantes),
        "errores": errores[:10],
    }


# ---------------------------------------------------------------------------
# Stats
# ---------------------------------------------------------------------------


def stats(db: Session) -> dict:
    docs_total = db.execute(
        select(func.count()).select_from(Item).where(Item.media_tipo == "documento")
    ).scalar() or 0
    docs_con_attachment = db.execute(
        select(func.count(func.distinct(Attachment.item_id))).where(Attachment.tipo == "documento")
    ).scalar() or 0
    docs_extraidos = db.execute(
        select(func.count()).select_from(Item).where(
            Item.media_tipo == "documento",
            Item.datos["extraccion"]["extracted_at"].astext.isnot(None),
        )
    ).scalar() or 0
    jobs_pendientes = db.execute(
        select(func.count()).select_from(Job).where(Job.tipo == "extract", Job.estado == "pendiente")
    ).scalar() or 0
    jobs_fallidos = db.execute(
        select(func.count()).select_from(Job).where(Job.tipo == "extract", Job.estado == "fallido")
    ).scalar() or 0
    return {
        "docs_total": int(docs_total),
        "docs_con_attachment": int(docs_con_attachment),
        "docs_extraidos": int(docs_extraidos),
        "docs_sin_binario": int(docs_total - docs_con_attachment),
        "jobs_extract_pendientes": int(jobs_pendientes),
        "jobs_extract_fallidos": int(jobs_fallidos),
    }
